from docx import Document
from docx.shared import Pt

doc = Document()

# Header
doc.add_heading('Raju Rastogi', 0)
doc.add_paragraph('New Delhi, India | raju.rastogi@ice.edu | +91-9876543210')
doc.add_paragraph('LinkedIn: linkedin.com/in/raju-rastogi-real')

# Profile
doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
    "Electrical Engineer with a transformed perspective on innovation and failure. "
    "Initially struggled with academic pressure but overcame personal challenges to develop "
    "a fearless approach to engineering. Passionate about honest work and practical problem-solving."
)

# Education
doc.add_heading('Education', level=1)
p = doc.add_paragraph()
p.add_run('Bachelor of Technology in Electrical Engineering').bold = True
p.add_run('\nImperial College of Engineering (ICE), Delhi').italic = True
p.add_run('\n2005 - 2009')
doc.add_paragraph('• CGPA: 5.5/10 (Consistently improved in final semesters after overcoming anxiety)')
doc.add_paragraph('• Key Subjects: Electrical Machines, Control Systems, Power Electronics')

# Skills
doc.add_heading('Skills', level=1)
doc.add_paragraph('• Engineering: Circuit Design, Soldering, Component Testing')
doc.add_paragraph('• Soft Skills: Radical Honesty, Resilience, Adaptability')
doc.add_paragraph('• Languages: Hindi (Native), English (Proficient)')

# Projects
doc.add_heading('Projects', level=1)
p = doc.add_paragraph()
p.add_run('Remote Controlled Drone (Team Project)').bold = True
doc.add_paragraph('• Assisted in the calibration of motor controls for a surveillance drone prototype.')
doc.add_paragraph('• Learned value of innovation over rote memorization.')

# Experience
doc.add_heading('Experience', level=1)
p = doc.add_paragraph()
p.add_run('Junior Engineer').bold = True
p.add_run('\nAnand Construction & Engineering').italic = True
p.add_run('\n2009 - Present')
doc.add_paragraph('• Managing site electrical layouts and safety compliances.')
doc.add_paragraph('• Applying practical engineering concepts to solve real-world site issues.')

doc.save('Raju_Rastogi_Resume.docx')
print("Resume created successfully: Raju_Rastogi_Resume.docx")
